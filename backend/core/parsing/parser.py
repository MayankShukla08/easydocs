from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParsedPage:
    page: Optional[int]
    text: str


def _extract_text_from_pypdf_page(page, page_num: int) -> str:
    """Extract text from a pypdf page using multiple strategies."""
    text_parts: list[str] = []

    # 1. Standard extraction
    try:
        txt = page.extract_text() or ""
        if txt.strip():
            text_parts.append(txt.strip())
    except Exception as e:
        logger.debug("pypdf standard extract failed on page %d: %s", page_num, e)

    # 2. Layout mode extraction if standard extraction produced little text
    if not text_parts or len("".join(text_parts)) < 30:
        try:
            txt_layout = page.extract_text(extraction_mode="layout") or ""
            if txt_layout.strip() and txt_layout.strip() not in text_parts:
                text_parts.append(txt_layout.strip())
        except Exception as e:
            logger.debug("pypdf layout extract failed on page %d: %s", page_num, e)

    # 3. Check for annotations / text comments
    try:
        if "/Annots" in page:
            annots = page.get("/Annots")
            if annots:
                for annot in annots:
                    annot_obj = annot.get_object() if hasattr(annot, "get_object") else annot
                    contents = annot_obj.get("/Contents") if hasattr(annot_obj, "get") else None
                    if contents and isinstance(contents, str) and contents.strip():
                        text_parts.append(contents.strip())
    except Exception as e:
        logger.debug("pypdf annotation extraction failed on page %d: %s", page_num, e)

    return "\n".join(text_parts).strip()


def parse_pdf(path: Path) -> list[ParsedPage]:
    pages: list[ParsedPage] = []

    # Try pypdf first (modern and feature-rich)
    try:
        import pypdf
        reader = pypdf.PdfReader(str(path))
        for i, p in enumerate(reader.pages, start=1):
            txt = _extract_text_from_pypdf_page(p, i)
            pages.append(ParsedPage(page=i, text=txt))

        total_extracted = sum(len(p.text) for p in pages)
        logger.info(
            "Parsed PDF %s with pypdf: %d pages, %d chars extracted",
            path.name,
            len(pages),
            total_extracted,
        )
        return pages
    except Exception as e:
        logger.warning("pypdf extraction failed for %s (%s), falling back to PyPDF2", path.name, e)

    # Fallback to PyPDF2
    try:
        from PyPDF2 import PdfReader as PyPDF2Reader
        reader = PyPDF2Reader(str(path))
        pages = []
        for i, p in enumerate(reader.pages, start=1):
            try:
                txt = p.extract_text() or ""
            except Exception as e:
                logger.warning("PyPDF2 page extract failed", extra={"page": i, "err": str(e)[:200]})
                txt = ""
            pages.append(ParsedPage(page=i, text=txt.strip()))

        total_extracted = sum(len(p.text) for p in pages)
        logger.info(
            "Parsed PDF %s with PyPDF2: %d pages, %d chars extracted",
            path.name,
            len(pages),
            total_extracted,
        )
        return pages
    except Exception as e:
        logger.error("All PDF parsers failed for %s: %s", path.name, e)
        raise


def parse_docx(path: Path) -> list[ParsedPage]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts: list[str] = []

    # 1. Paragraphs
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())

    # 2. Tables (crucial for documents formatted in tables / calendars / schedules)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_cells:
                # Deduplicate identical consecutive cells from merged columns
                deduped: list[str] = []
                for c in row_cells:
                    if not deduped or deduped[-1] != c:
                        deduped.append(c)
                parts.append(" | ".join(deduped))

    combined = "\n".join(parts)
    logger.info("Parsed DOCX %s: %d parts, %d chars extracted", path.name, len(parts), len(combined))
    return [ParsedPage(page=1, text=combined)]


def parse_file(path: Path, mime_type: str) -> list[ParsedPage]:
    mt = (mime_type or "").lower()
    if mt == "application/pdf" or path.suffix.lower() == ".pdf":
        return parse_pdf(path)
    if mt in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or path.suffix.lower() == ".docx":
        return parse_docx(path)
    raise ValueError(f"Unsupported file type: {mime_type}")


